

# Imports
from langgraph.graph import START, END, StateGraph, MessagesState
from langgraph.checkpoint.memory import MemorySaver
from langgraph.prebuilt import ToolNode
from langchain.messages import HumanMessage, AIMessage, SystemMessage
from langchain_core.tools import tool
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
# from langchain.vectorstores import FAISS
from langchain_chroma import Chroma
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, PyPDFDirectoryLoader

from dotenv import load_dotenv
# from IPython.display import Image, display
from typing import Literal
import os

print(" All imports successful")

# Load API key
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

if not openai_api_key:
    raise ValueError("OPENAI_API_KEY not found! Please set it in your .env file.")

print("API key loaded")

# Initialize LLM
llm = ChatOpenAI(
    model="gpt-5-nano",
    temperature=0.5,
    api_key=openai_api_key
)

print(f"LLM initialized: {llm.model_name}")

# Load all PDFs, DOCX and TXT files from the Tax_Project folder
# Optional: python-docx may be required to read .docx files
try:
    from docx import Document as DocxReader
    _HAS_DOCX = True
except Exception:
    _HAS_DOCX = False

folder = 'TAX_Documents'
if not os.path.isdir(folder):
    raise FileNotFoundError(f"Folder not found: {folder}. Update path accordingly.")

documents = []

for fname in sorted(os.listdir(folder)):
    fpath = os.path.join(folder, fname)
    if os.path.isdir(fpath):
        continue
    lower = fname.lower()
    try:
        if lower.endswith('.pdf'):
            loader = PyPDFLoader(fpath)
            # load() returns a list of Document objects (pages)
            pages = loader.load()
            for p in pages:
                # add filename to metadata so you can trace source
                p.metadata = {**getattr(p, 'metadata', {}), 'source': fname}
                documents.append(p)
        elif lower.endswith('.txt'):
            with open(fpath, 'r', encoding='utf-8') as f:
                text = f.read()
            documents.append(Document(page_content=text, metadata={'source': fname}))
        elif lower.endswith('.docx'):
            if not _HAS_DOCX:
                raise ImportError('python-docx is not installed. Run: python -m pip install python-docx')
            doc = DocxReader(fpath)
            text = '\n'.join(p.text for p in doc.paragraphs)
            documents.append(Document(page_content=text, metadata={'source': fname}))
        else:
            # skip other file types
            continue
    except Exception as e:
        print(f"Failed to load {fname}: {e}")

print(f"Loaded {len(documents)} documents from {folder}")

# Create text splitter
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,      # Characters per chunk
    chunk_overlap=100     # Overlap to preserve context
)

# Split documents
doc_splits = text_splitter.split_documents(pages)

print(f" Created {len(doc_splits)} chunks")
print(f"\nSample chunk:")
print(f"{doc_splits[0].page_content[:200]}...")


# Initialize embeddings (using OpenAI)
embeddings = OpenAIEmbeddings(
    model="text-embedding-3-small",
    api_key=openai_api_key
)

print(" Embeddings model initialized")

# Create FAISS vector store
faiss_path = "./faiss_index_agentic_rag"

# Create vector store from documents (build in-memory FAISS index)
vectorstore = FAISS.from_documents(documents=doc_splits, embedding=embeddings)

# Persist FAISS index locally
vectorstore.save_local(faiss_path)

print(f" FAISS vector store created with {len(doc_splits)} chunks")
print(f"   Persisted to: {faiss_path}")


# defining tool
@tool
def retrieve_documents(query: str) -> str:
    """
    Search for relevant documents in the knowledge base.
    
    Use this tool when you need information from the document collection
    to answer the user's question. Do NOT use this for:
    - General knowledge questions
    - Greetings or small talk
    - Simple calculations
    
    Args:
        query: The search query describing what information is needed
        
    Returns:
        Relevant document excerpts that can help answer the question
    """
    # Use MMR (Maximum Marginal Relevance) for diverse results
    retriever = vectorstore.as_retriever(
        search_type="mmr",
        search_kwargs={"k": 5, "fetch_k": 10}
    )
    
    # Retrieve documents
    results = retriever.invoke(query)
    
    if not results:
        return "No relevant documents found."
    
    # Format results
    formatted = "\n\n---\n\n".join(
        f"Document {i+1}:\n{doc.page_content}"
        for i, doc in enumerate(results)
    )
    
    return formatted

print(" Retrieval tool created")


#System Prompt
system_prompt = SystemMessage(content="""You are a helpful assistant with access to a document retrieval tool. Prioritize documents in the Tax_Project folder for answers and cite them.

RETRIEVAL DECISION RULES:

DO NOT retrieve for:
- Greetings, capability questions (e.g. "What can you help with?"), simple math or general knowledge, or casual conversation.

DO retrieve for:
- Questions asking for specific information that would be in documents.
- Requests for facts, definitions, or explanations about specialized topics (tax policy, statutes, guidance).
- Any question where citing sources would improve the answer.

ADDITIONAL RULES:
- Jurisdiction: ask the user if not specified; prefer documents applicable to the stated jurisdiction.
- Recency: prefer documents published after 2015 unless the user requests otherwise.
- Retrieval limits: retrieve at most 5 documents and include up to 300 characters per excerpt.
- Citation format: append " — Source: filename (page N)" to any excerpt or claim derived from a document.
- Conflicts: if documents disagree, list each source and explicitly state which you prefer and why.
- Clarify: if the query is ambiguous, ask one clarifying question before retrieving.
- Privacy: do not reveal private or sensitive content from documents unless the user explicitly permits it.

When you retrieve documents, cite them in your answer. If documents do not contain the answer, say so.
""")

print(" System prompt configured")

# Bind tools to LLM
tools = [retrieve_documents, tax_calculator]
llm_with_tools = llm.bind_tools(tools)

def assistant(state: MessagesState) -> dict:
    """Assistant node - decides whether to retrieve or answer directly."""
    messages = [system_prompt] + state['messages']
    response = llm_with_tools.invoke(messages)
    return {'messages': [response]}

def should_continue(state: MessagesState) -> Literal['tools', '__end__']:
    """Decide whether to call tools or finish."""
    last_message = state['messages'][-1]
    if last_message.tool_calls:
        return 'tools'
    return '__end__'

print(" Agent nodes defined")

# Build graph
builder = StateGraph(MessagesState)

# Add nodes
builder.add_node('assistant', assistant)
builder.add_node('tools', ToolNode(tools))

# Define edges
builder.add_edge(START, 'assistant')
builder.add_conditional_edges(
    'assistant',
    should_continue,
    {'tools': 'tools', '__end__': END},
)
builder.add_edge('tools', 'assistant')

# Add memory
memory = MemorySaver()
agent = builder.compile(checkpointer=memory)

print(" Agentic RAG system compiled")
